import json
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.table import WD_ALIGN_VERTICAL, WD_ROW_HEIGHT_RULE
import os
import io

from utils.xml_helpers import *
# === CONFIGURATION ===
HEADER_IMAGE = 'logo.png'
OUTPUT_FILENAME = "Generated_Resume.docx"

# Colors
BLUE = RGBColor(0, 102, 204)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
BLACK = RGBColor(0, 0, 0)


def resume_builder(json_data):
    doc = Document()
    for section in doc.sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.25)
        section.right_margin = Inches(0.25)

    # FIRST PAGE: 3-column layout (Sidebar | Spacer | Main Content)
    main_table = doc.add_table(rows=2, cols=3)
    main_table.autofit = False

    # Sidebar spans both rows
    sidebar_cell = main_table.cell(0, 0)
    sidebar_cell.merge(main_table.cell(1, 0))

    # Spacer spans both rows
    spacer_cell = main_table.cell(0, 1)
    spacer_cell.merge(main_table.cell(1, 1))

    # Main content split into profile and experience
    profile_container_cell = main_table.cell(0, 2)
    experience_container_cell = main_table.cell(1, 2)

    # Column widths
    sidebar_cell.width = Inches(2.4)
    spacer_cell.width = Inches(0.15)
    # ### CHANGE 1: Reduced width slightly to prevent right border cut-off ###
    profile_container_cell.width = Inches(5.2)
    experience_container_cell.width = Inches(5.2)


    # Remove borders in spacer
    tc_borders = OxmlElement('w:tcBorders')
    spacer_cell._tc.get_or_add_tcPr().append(tc_borders)

    # Sidebar content
    logo_p = sidebar_cell.add_paragraph()
    logo_p.paragraph_format.space_before = Pt(0)
    logo_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    if os.path.exists(HEADER_IMAGE):
        logo_p.add_run().add_picture(HEADER_IMAGE, width=Inches(2.4))
    else:
        logo_p.add_run("KANERIKA").font.size = Pt(14)

    blue_sidebar_table = sidebar_cell.add_table(rows=1, cols=1)
    blue_sidebar_cell = blue_sidebar_table.cell(0, 0)
    set_cell_background(blue_sidebar_cell, BLUE)
    blue_sidebar_cell.width = Inches(2.4)

    # ... (rest of the sidebar code is unchanged) ...
    name_para = blue_sidebar_cell.add_paragraph()
    name_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    name_run = name_para.add_run(json_data.get("name", "").upper())
    name_run.font.name = 'Aptos'
    name_run.font.color.rgb = WHITE
    name_run.font.size = Pt(14)
    name_run.bold = True
    name_run.underline = True
    name_para.paragraph_format.space_after = Pt(6)

    contact_info = json_data.get("contact", {})
    email_para = blue_sidebar_cell.add_paragraph()
    email_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    email_para.add_run('✉️ ').font.color.rgb = WHITE
    email_run = email_para.add_run(contact_info.get("email", ""))
    email_run.font.name = 'Aptos'
    email_run.font.color.rgb = WHITE
    email_run.font.size = Pt(9)
    email_para.paragraph_format.space_after = Pt(20)

    add_sidebar_separator(blue_sidebar_cell)

    add_header(blue_sidebar_cell, "EDUCATION", WHITE, 12, font_name='Aptos', align=WD_ALIGN_PARAGRAPH.CENTER)
    for edu in json_data.get("education", []):
        edu_para = blue_sidebar_cell.add_paragraph()
        edu_para.add_run(f'{edu.get("degree", "")}\n').font.color.rgb = WHITE
        edu_para.add_run(f'{edu.get("institution", "")}\n').font.color.rgb = WHITE
        edu_para.add_run(f'{edu.get("location", "")}\n').font.color.rgb = WHITE
        edu_para.add_run(f'{edu.get("start_date", "")} - {edu.get("end_date", "")}').font.color.rgb = WHITE
        for run in edu_para.runs:
            run.font.name = 'Aptos'
            run.font.size = Pt(10)
        edu_para.paragraph_format.space_after = Pt(12)

    add_sidebar_separator(blue_sidebar_cell)

    add_header(blue_sidebar_cell, "SKILLS", WHITE, 12, font_name='Aptos', align=WD_ALIGN_PARAGRAPH.CENTER)
    skills_data = json_data.get("skills", {})
    all_skills = skills_data.get("technical", []) + skills_data.get("tools", [])
    if all_skills:
        for skill in all_skills[:10]:
            ps = blue_sidebar_cell.add_paragraph(skill, style='List Bullet')
            ps.runs[0].font.name = 'Aptos'
            ps.runs[0].font.size = Pt(10)
            ps.runs[0].font.color.rgb = WHITE
            ps.paragraph_format.left_indent = Inches(0.45)
            ps.paragraph_format.space_after = Pt(2)
            ps.paragraph_format.space_before = Pt(2)

    blue_sidebar_cell.add_paragraph().paragraph_format.space_after = Pt(12)

    add_sidebar_separator(blue_sidebar_cell)

    if json_data.get("certifications"):
        add_header(blue_sidebar_cell, "CERTIFICATIONS", WHITE, 12, font_name='Aptos', align=WD_ALIGN_PARAGRAPH.CENTER)
        for cert in json_data.get("certifications"):
            pc = blue_sidebar_cell.add_paragraph(cert)
            pc.alignment = WD_ALIGN_PARAGRAPH.LEFT
            pc.runs[0].font.name = 'Aptos'
            pc.runs[0].font.size = Pt(9)
            pc.runs[0].font.color.rgb = WHITE
            pc.paragraph_format.space_after = Pt(1)

    # Main content: Profile
    add_header(profile_container_cell, "PROFILE", BLUE, 22, font_name='Aptos')
    profile_text = json_data.get("profile") or json_data.get("summary")
    if not profile_text:
        profile_text = "No profile information provided."
    if profile_text:
        profile_content_table = profile_container_cell.add_table(rows=1, cols=1)
        profile_content_table.autofit = False
        profile_content_table.columns[0].width = profile_container_cell.width
        profile_content_cell = profile_content_table.cell(0, 0)
        set_cell_borders(profile_content_cell, BLACK)
        set_cell_margins(profile_content_cell, top=5, start=5, bottom=5, end=5)
        p = profile_content_cell.add_paragraph()
        p.add_run(profile_text).font.size = Pt(11)
        p.runs[0].font.name = 'Aptos'
        p.paragraph_format.space_after = Pt(0)

    ### CHANGE 2: Set the height for the second row to expand the container ###
    experience_row = main_table.rows[1]
    experience_row.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
    experience_row.height = Inches(8)

    # Main content: Professional Experience
    add_header(experience_container_cell, "PROFESSIONAL EXPERIENCE", BLUE, 22, font_name='Aptos')
    experience_list = json_data.get("experience", [])
    first_page_exp = experience_list[:1]
    if first_page_exp:
        # Create a nested table for the bordered content
        exp_content_table = experience_container_cell.add_table(rows=1, cols=1)
        
        exp_content_table.autofit = False
        exp_content_table.columns[0].width = experience_container_cell.width

        # ### CHANGE 3: Make the nested table's row fill the new container height ###
        exp_content_row = exp_content_table.rows[0]
        exp_content_row.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
        exp_content_row.height = experience_row.height
        
        exp_content_cell = exp_content_table.cell(0, 0)
        exp_content_cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
        set_cell_borders(exp_content_cell, BLACK)
        set_cell_margins(exp_content_cell, top=0, start=5, bottom=0, end=0)
        add_experience_entry(exp_content_cell, first_page_exp[0])

    # Second page onwards
    remaining_exp = experience_list[1:]
    if remaining_exp:
        doc.add_page_break()
        add_header(doc, "PROFESSIONAL EXPERIENCE", BLUE, 22, font_name='Aptos')

        page2_table = doc.add_table(rows=1, cols=1)
        page2_table.autofit = False
        section = doc.sections[-1]
        section.left_margin = Inches(0.35) 
        usable_width = section.page_width - section.left_margin - section.right_margin
        page2_table.columns[0].width = usable_width
        
        # ### CHANGE 1: Calculate the available height on the page ###
        # We subtract an estimated height for the header above. You can adjust Inches(0.8).
        usable_height = section.page_height - section.top_margin - section.bottom_margin
        table_height = usable_height - Inches(0.8) 
        
        # ### CHANGE 2: Apply the calculated height to the table row ###
        row = page2_table.rows[0]
        row.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
        row.height = table_height

        container_cell = page2_table.cell(0, 0)
        # ### CHANGE 3: Align content to the top of the now-tall cell ###
        container_cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
        
        set_cell_borders(container_cell, BLACK)
        set_cell_margins(container_cell, top=5, start=8, bottom=5, end=8)

        for exp in remaining_exp:
            add_experience_entry(container_cell, exp)

     # Save to BytesIO
    file_stream = io.BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)  # Move pointer to the beginning
    
    print(f"Resume document '{OUTPUT_FILENAME}' has been created.")
    return file_stream.getvalue()

if __name__ == '__main__':
    try:
        with open('resume_data.json', 'r') as f:
            json_data = json.load(f)
        resume_builder(json_data)
    except FileNotFoundError:
        print("Error: 'resume_data.json' not found.")
    except json.JSONDecodeError:
        print("Error: 'resume_data.json' is not valid JSON.")
