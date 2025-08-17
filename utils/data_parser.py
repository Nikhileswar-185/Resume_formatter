"""
This module provides functions to extract text from PDF and DOCX files in memory.
It supports both plain text extraction and basic table cell content.
"""


import fitz  # PyMuPDF
from docx import Document
from io import BytesIO

def extract_text_from_pdf(file_bytes: bytes) -> str:
    """Extract plain text from a PDF file in memory.

    Parameters
    - file_bytes: Raw PDF bytes

    Returns
    - str: Combined text content for all pages
    """
    text_parts = []
    with fitz.open(stream=file_bytes, filetype="pdf") as doc:
        for page in doc:
            text_parts.append(page.get_text("text"))
    return "\n".join(text_parts)

def extract_text_from_docx(file_bytes: bytes) -> str:
    """Extract text from a DOCX file in memory.

    Parameters
    - file_bytes: Raw docx bytes

    Returns
    - str: Combined text content from paragraphs and basic table cells
    """
    buffer = BytesIO(file_bytes)
    doc = Document(buffer)
    texts = []
    for para in doc.paragraphs:

        if para.text and para.text.strip():
            texts.append(para.text.strip())
    # Tables (optional)
    for table in doc.tables:
        for row in table.rows:
            row_text = [cell.text.strip() for cell in row.cells if cell.text and cell.text.strip()]
            if row_text:
                texts.append(" | ".join(row_text))
    return "\n".join(texts)


def extract_text_from_file(file_bytes: bytes, ext: str) -> str:
    """Route text extraction based on file extension.

    Parameters
    - file_bytes: Original file bytes
    - ext: Lowercased extension without dot (e.g., "pdf", "docx")

    Returns
    - str: Extracted text

    Raises
    - ValueError: For unsupported file types
    """
    if ext.lower() == "pdf":
        return extract_text_from_pdf(file_bytes)
    elif ext.lower() == "docx":
        return extract_text_from_docx(file_bytes)
    else:
        raise ValueError(f"Unsupported file type: {ext}")