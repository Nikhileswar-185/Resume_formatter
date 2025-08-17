"""
This module contains utility functions for manipulating XML elements in Word documents.
It provides functions to set cell background colors, borders, margins, and add headers and bullet points
"""


import json
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


BLUE = RGBColor(0, 102, 204)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
BLACK = RGBColor(0, 0, 0)

def set_cell_background(cell, color):
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:val'), 'clear')
    shading_elm.set(qn('w:fill'), f'{color[0]:02x}{color[1]:02x}{color[2]:02x}')
    cell._tc.get_or_add_tcPr().append(shading_elm)

def set_cell_borders(cell, color, size_pt=1):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_borders = OxmlElement('w:tcBorders')
    for border_name in ('top', 'bottom', 'left', 'right'):
        border_el = OxmlElement(f'w:{border_name}')
        border_el.set(qn('w:val'), 'single')
        border_el.set(qn('w:sz'), str(int(size_pt * 8)))
        border_el.set(qn('w:color'), f'{color[0]:02x}{color[1]:02x}{color[2]:02x}')
        tc_borders.append(border_el)
    tc_pr.append(tc_borders)

def set_cell_margins(cell, top=0, start=0, bottom=0, end=0):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, v in [("top", top), ("start", start), ("bottom", bottom), ("end", end)]:
        mar = OxmlElement(f'w:{m}')
        mar.set(qn('w:w'), str(v * 20))
        mar.set(qn('w:type'), 'dxa')
        tcMar.append(mar)
    tcPr.append(tcMar)

def add_header(parent_obj, text, color, size, is_bold=True, align=WD_ALIGN_PARAGRAPH.LEFT, font_name='Aptos'):
    header = parent_obj.add_paragraph()
    run = header.add_run(text.upper())
    run.font.color.rgb = color
    run.font.size = Pt(size)
    run.bold = is_bold
    run.font.name = font_name
    run.underline = True
    header.alignment = align
    header.paragraph_format.space_after = Pt(2)
    return header

def add_bullet_points(parent_obj, items, color=None, font_name='Roboto'):
    for item in items:
        p = parent_obj.add_paragraph(style='List Bullet')
        run = p.add_run(item.replace('\n', ' '))
        run.font.name = font_name
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.left_indent = Inches(0.75)
        if color:
            run.font.color.rgb = color

def add_experience_entry(parent_obj, exp):
    exp_table = parent_obj.add_table(rows=1, cols=2)
    exp_table.autofit = True

    title_cell = exp_table.cell(0, 0)
    date_cell = exp_table.cell(0, 1)
    title_cell.width = Inches(3.5)
    date_cell.width = Inches(2.5)

    title_p = title_cell.add_paragraph()
    title_run = title_p.add_run(f'{exp.get("title", "")}')
    title_p.paragraph_format.left_indent = Pt(0)
    title_p.paragraph_format.first_line_indent = Pt(0)
    title_run.bold = True
    title_run.font.size = Pt(11)
    title_run.font.name = 'Aptos'
    title_run.italic = True

    company_run = title_p.add_run(f' at {exp.get("company", "")}')
    company_run.font.size = Pt(10)
    company_run.font.name = 'Aptos'
    company_run.italic = True
    company_run.bold = True
    # title_p.paragraph_format.space_after = Pt(0)

    date_p = date_cell.add_paragraph()
    date_run = date_p.add_run(f'({exp.get("start_date", "")} - {exp.get("end_date", "")})')
    date_run.font.size = Pt(11)
    date_run.font.name = 'Aptos'
    date_run.bold = True
    date_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    # date_p.paragraph_format.space_after = Pt(0)

    add_bullet_points(parent_obj, exp.get("achievements", []))
    # parent_obj.add_paragraph().paragraph_format.space_after = Pt(12)

def add_sidebar_separator(cell, width_ratio=3):
    p = cell.add_paragraph()
    total_spaces = int(40 * width_ratio)
    line_run = p.add_run("_" * total_spaces)
    line_run.font.size = Pt(3.5)
    line_run.font.color.rgb = RGBColor(255, 255, 255)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    line_run.bold = True
    p.paragraph_format.right_indent = Inches(0.1)